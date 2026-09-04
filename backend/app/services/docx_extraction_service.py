# services/docx_extraction_service.py
import docx
from docx.oxml.ns import qn


class DocxExtractionService:

    def extract_text(self, docx_path: str) -> str:
        document = docx.Document(docx_path)
        parts = []

        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts)

    def extract_hyperlinks(self, docx_path: str) -> dict:
        document = docx.Document(docx_path)
        rels = document.part.rels

        result = {
            "github_url": None,
            "linkedin_url": None,
            "portfolio_url": None,
            "website_url": None,
        }

        for rel in rels.values():
            if rel.reltype == docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK and rel.target_ref:
                url = rel.target_ref
                lower = url.lower()
                if "github.com" in lower and not result["github_url"]:
                    result["github_url"] = url
                elif "linkedin.com" in lower and not result["linkedin_url"]:
                    result["linkedin_url"] = url
                elif not result["website_url"]:
                    result["website_url"] = url

        return result