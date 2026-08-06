import io

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

from app.schemas.resume import ResumeDocument


def add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run_el.append(rpr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run_el.append(text_el)

    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)

    return hyperlink


def generate_resume_docx(resume: ResumeDocument) -> bytes:

    document = docx.Document()
    document.styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = 1
    run = title.add_run(resume.name)
    run.bold = True
    run.font.size = Pt(18)

    headline = document.add_paragraph(resume.headline)
    headline.alignment = 1

    contact = document.add_paragraph()
    contact.alignment = 1
    c = resume.contact_info

    if c.location:
        contact.add_run(c.location + "  |  ")
    if c.phone:
        contact.add_run(c.phone + "  |  ")
    if c.email:
        contact.add_run(c.email)
    if c.github:
        contact.add_run("  |  ")
        add_hyperlink(contact, c.github_url, c.github) if c.github_url else contact.add_run(c.github)
    if c.linkedin:
        contact.add_run("  |  ")
        add_hyperlink(contact, c.linkedin_url, c.linkedin) if c.linkedin_url else contact.add_run(c.linkedin)
    if c.portfolio:
        contact.add_run("  |  ")
        add_hyperlink(contact, c.portfolio_url, c.portfolio) if c.portfolio_url else contact.add_run(c.portfolio)

    def add_heading(text):
        h = document.add_heading(text.upper(), level=2)
        h.runs[0].font.size = Pt(12)

    add_heading("Professional Summary")
    document.add_paragraph(resume.professional_summary.content)

    add_heading("Technical Skills")
    for category in resume.technical_skills.categories:
        p = document.add_paragraph()
        p.add_run(f"{category.category}: ").bold = True
        p.add_run(", ".join(category.skills))

    add_heading("Professional Experience")
    for exp in resume.professional_experience:
        p = document.add_paragraph()
        p.add_run(exp.role).bold = True
        p.add_run(f"    {exp.start_date} - {exp.end_date}")

        sub = exp.company + (f", {exp.location}" if exp.location else "")
        document.add_paragraph(sub)

        for point in exp.responsibilities:
            document.add_paragraph(point, style="List Bullet")

        for project in exp.projects:
            proj_p = document.add_paragraph()
            proj_p.add_run(project.title).bold = True
            for bullet in project.bullet_points:
                document.add_paragraph(bullet, style="List Bullet")

    if resume.certifications:
        add_heading("Certifications")
        for cert in resume.certifications:
            document.add_paragraph(cert.name, style="List Bullet")

    if resume.education:
        add_heading("Education")
        for edu in resume.education:
            line = f"{edu.degree} — {edu.institution}"
            if edu.start_year or edu.end_year:
                line += f" ({edu.start_year} - {edu.end_year})"
            document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()